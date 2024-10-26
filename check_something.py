from docx import Document
import json
import os

# Загрузка JSON-файлов
json_files = [f for f in os.listdir("Sortings") if f.endswith('.json')]

for key in json_files:
    with open(f'Sortings/{key}', 'r') as f:
        data = json.load(f)

    # Создание документа Word
    doc = Document()
    doc.add_heading(f'{key} Timing Table', level=1)

    # Определение всех заголовков (Case types) для первой строки
    case_types = list(data.keys())
    # Получение списка значений N (из первого тестового случая, например, "Sorted")
    n_values = list(data[case_types[0]].keys())

    # Создание таблицы с числом строк равным числу элементов в n_values и числом столбцов по количеству Case types + 1 для N
    table = doc.add_table(rows=len(n_values) + 1, cols=len(case_types) + 1)
    table.style = 'Table Grid'

    # Заполнение заголовков столбцов (Case types) в первой строке
    header_cells = table.rows[0].cells
    header_cells[0].text = 'N'
    for i, case in enumerate(case_types):
        header_cells[i + 1].text = case

    # Заполнение строк значениями N и времени выполнения для каждого случая
    for i, n in enumerate(n_values):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = n  # Значение N
        for j, case in enumerate(case_types):
            # Используем get, чтобы избежать KeyError, если значение N отсутствует для данного case
            row_cells[j + 1].text = str(data[case].get(n, "N/A"))

    # Сохранение документа
    os.makedirs("Tables", exist_ok=True)  # Создаем папку, если ее нет
    doc.save(f'Tables/{key}_Timings.docx')
    print(f'Создан документ: Tables/{key}_Timings.docx')
